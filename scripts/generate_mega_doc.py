"""generate_mega_doc.py -- Gabungkan seluruh docs/ menjadi satu dokumen.

Menghasilkan:
    docs/DOKUMENTASI-HANDOVER-LENGKAP.md    (selalu)
    docs/DOKUMENTASI-HANDOVER-LENGKAP.docx  (dengan --docx; butuh pypandoc-binary + python-docx)

Link antar-dokumen dikonversi menjadi link internal (#sec-...), dan LABEL link
yang berupa nama file diganti judul section-nya supaya tidak ada teks yang
terbaca merujuk ke file .md lain.

Pakai:
    python scripts/generate_mega_doc.py            # md saja
    python scripts/generate_mega_doc.py --docx     # md + docx cantik
"""

import argparse
import os
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUT_MD = DOCS / "DOKUMENTASI-HANDOVER-LENGKAP.md"
OUT_DOCX = DOCS / "DOKUMENTASI-HANDOVER-LENGKAP.docx"

TITLE = "Dashboard SPEED Pertamina Energy Institute"
SUBTITLE = "Dokumentasi Lengkap — Referensi Teknis, How-To, dan Paket Serah Terima"

ORDER = [
    ("Referensi Teknis", "01-arsitektur.md"),
    ("Referensi Teknis", "02-migrasi-storage.md"),
    ("Referensi Teknis", "03-database.md"),
    ("Referensi Teknis", "04-pipeline-scheduling.md"),
    ("Referensi Teknis", "05-sumber-data.md"),
    ("Referensi Teknis", "06-ai-sentiment.md"),
    ("Referensi Teknis", "07-power-bi.md"),
    ("Referensi Teknis", "08-maintenance.md"),
    ("Referensi Teknis", "09-pengembangan.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/README.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/01-setup-lokal-dari-nol.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/02-menjalankan-pipeline-manual.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/03-cek-kesehatan-scheduler.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/04-backfill-data-bolong.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/05-menambah-topik-berita.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/06-menambah-sumber-terstruktur.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/07-koneksi-power-bi-neon.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/08-rotasi-kredensial.md"),
    ("How-To (Langkah-demi-Langkah)", "how-to/09-backup-restore-neon.md"),
    ("Paket Serah Terima (Handover)", "handover/01-bast.md"),
    ("Paket Serah Terima (Handover)", "handover/02-inventaris-aset-akses.md"),
    ("Paket Serah Terima (Handover)", "handover/03-runbook-hari-pertama.md"),
    ("Paket Serah Terima (Handover)", "handover/04-biaya-lisensi.md"),
    ("Paket Serah Terima (Handover)", "handover/05-diagram-alur-data.md"),
]


def doc_anchor(rel: str) -> str:
    return "sec-" + re.sub(r"[^a-z0-9]+", "-", rel.lower()).strip("-")


def first_heading(text: str) -> str | None:
    for ln in text.splitlines():
        m = re.match(r"^#{1,6} (.+)$", ln)
        if m:
            return m.group(1).strip()
    return None


def looks_like_filename(label: str, base: str) -> bool:
    """Label link yang berupa nama file/path, bukan kalimat manusiawi."""
    l = label.strip()
    return (
        l == base
        or l.endswith(".md")
        or ("/" in l and " " not in l)
    )


def build_mega() -> str:
    name_to_anchor: dict[str, str] = {}
    name_to_title: dict[str, str] = {}
    texts: dict[str, str] = {}

    for _, rel in ORDER:
        text = (DOCS / rel).read_text(encoding="utf-8").rstrip() + "\n"
        texts[rel] = text
        base = os.path.basename(rel)
        name_to_anchor[base] = doc_anchor(rel)
        name_to_title[base] = first_heading(text) or rel

    def transform_links(text: str, src_rel: str) -> str:
        src_dir = (DOCS / src_rel).parent

        def repl(m):
            label, target = m.group(1), m.group(2).strip()
            if target.startswith(("http://", "https://", "mailto:", "#", "<")):
                return m.group(0)
            path = target.partition("#")[0]
            base = os.path.basename(path)
            if base in name_to_anchor and base.endswith(".md"):
                if looks_like_filename(label, base):
                    label = f"bagian “{name_to_title[base]}”"
                return f"[{label}](#{name_to_anchor[base]})"
            absolute = os.path.normpath(str(src_dir / path))
            newrel = os.path.relpath(absolute, DOCS).replace("\\", "/")
            return f"[{label}]({newrel})"

        return re.sub(r"\[([^\]]*)\]\(([^)]+)\)", repl, text)

    sections, toc, group_now = [], [], None
    for group, rel in ORDER:
        text = texts[rel]
        title = first_heading(text) or rel
        anchor = doc_anchor(rel)
        if group != group_now:
            toc.append(f"\n**{group}**\n")
            group_now = group
        toc.append(f"- [{title}](#{anchor})")
        sections.append(f'<a id="{anchor}"></a>\n\n{transform_links(text, rel)}')

    header = (
        f"#### Dokumentasi Lengkap — {TITLE}\n\n"
        "> Dokumen gabungan (auto-generated oleh `scripts/generate_mega_doc.py`) dari "
        f"{len(ORDER)} file di `docs/`, `docs/how-to/`, `docs/handover/`. "
        "Sumber kebenaran tetap file per-topik; jalankan ulang script bila ada perubahan.\n\n"
        "##### Daftar Isi\n"
    )
    return header + "\n".join(toc) + "\n\n---\n\n" + "\n\n---\n\n".join(sections)


# ── Rendering diagram ─────────────────────────────────────────────────────────

ASSETS = DOCS / "assets"
_BOX_CHARS = "─│┌┐└┘├┤┬┴▼▲◄►═║╔╗╚╝"


def render_mermaid(code: str) -> Path | None:
    """Render mermaid via kroki.io -> PNG (cache per-hash di docs/assets)."""
    import base64
    import hashlib
    import urllib.request
    import zlib

    ASSETS.mkdir(exist_ok=True)
    out = ASSETS / f"mermaid-{hashlib.sha1(code.encode()).hexdigest()[:12]}.png"
    if out.exists():
        return out
    enc = base64.urlsafe_b64encode(zlib.compress(code.encode(), 9)).decode()
    req = urllib.request.Request(f"https://kroki.io/mermaid/png/{enc}",
                                 headers={"User-Agent": "Mozilla/5.0"})
    try:
        out.write_bytes(urllib.request.urlopen(req, timeout=120).read())
        return out
    except Exception as exc:
        print(f"  [mermaid] gagal render (blok dibiarkan sebagai kode): {exc}")
        return None


def render_ascii_art(code: str) -> Path | None:
    """Render diagram ASCII (box-drawing) -> PNG monospace via Pillow."""
    import hashlib

    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return None

    ASSETS.mkdir(exist_ok=True)
    out = ASSETS / f"ascii-{hashlib.sha1(code.encode()).hexdigest()[:12]}.png"
    if out.exists():
        return out
    lines = code.rstrip("\n").splitlines()
    try:
        font = ImageFont.truetype("consola.ttf", 28)
    except OSError:
        font = ImageFont.load_default()
    bbox = font.getbbox("M")
    cw, ch = bbox[2] - bbox[0], int((bbox[3] - bbox[1]) * 1.6)
    W = max(len(l) for l in lines) * cw + 80
    H = len(lines) * ch + 80
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    for i, ln in enumerate(lines):
        draw.text((40, 40 + i * ch), ln, fill=(0x2B, 0x33, 0x3E), font=font)
    img.save(out)
    return out


def replace_diagram_blocks(text: str) -> str:
    """```mermaid``` -> gambar kroki; code block berisi box-drawing -> gambar PIL."""
    def sub_mermaid(m):
        p = render_mermaid(m.group(1))
        return f"![Diagram](assets/{p.name})" if p else m.group(0)

    text = re.sub(r"```mermaid\n(.*?)```", sub_mermaid, text, flags=re.S)

    def sub_ascii(m):
        body = m.group(2)
        if any(c in body for c in _BOX_CHARS):
            p = render_ascii_art(body)
            if p:
                return f"![Diagram](assets/{p.name})"
        return m.group(0)

    return re.sub(r"```([a-zA-Z]*)\n(.*?)```", sub_ascii, text, flags=re.S)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def md_to_docx(mega: str) -> None:
    import pypandoc

    text = re.sub(r"^#### Dokumentasi Lengkap[^\n]*\n", "", mega)
    text = re.sub(r"##### Daftar Isi\n.*?\n---\n", "", text, flags=re.S)
    text = re.sub(r'<a id="(sec-[^"]+)"></a>\s*\n+(#{1,6}) ([^\n]+)',
                  lambda m: f"{m.group(2)} {m.group(3)} {{#{m.group(1)}}}", text)
    text = replace_diagram_blocks(text)

    out, fence = [], False
    for ln in text.splitlines(keepends=True):
        if re.match(r"^\s*(```|~~~)", ln):
            fence = not fence
            out.append(ln)
            continue
        m = None if fence else re.match(r"^(#{1,6}) (.*)$", ln)
        out.append("#" * max(1, len(m.group(1)) - 3) + " " + m.group(2) + "\n" if m else ln)
    tmp = DOCS / "_mega_tmp.md"
    tmp.write_text("".join(out), encoding="utf-8", newline="\n")

    try:
        pypandoc.convert_file(
            str(tmp), "docx", format="markdown+emoji", outputfile=str(OUT_DOCX),
            extra_args=[
                f"--resource-path={DOCS}",
                "--toc", "--toc-depth=2",
                "--metadata", f"title={TITLE}",
                "--metadata", f"subtitle={SUBTITLE}",
                "--metadata", "date=Juli 2026",
                "--metadata", "lang=id",
                "--metadata", "toc-title=Daftar Isi",
            ],
        )
    finally:
        tmp.unlink(missing_ok=True)
    style_docx()


def style_docx() -> None:
    import docx
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document(str(OUT_DOCX))
    BLUE = RGBColor(0x00, 0x51, 0x9C)
    RED = RGBColor(0xDA, 0x29, 0x1C)
    GRAY = RGBColor(0x44, 0x4B, 0x55)

    def style_font(name, size=None, color=None, bold=None, font_name=None):
        try:
            st = d.styles[name]
        except KeyError:
            return
        if size: st.font.size = Pt(size)
        if color is not None: st.font.color.rgb = color
        if bold is not None: st.font.bold = bold
        if font_name:
            st.font.name = font_name
            st.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    style_font("Title", 26, BLUE, True, "Calibri Light")
    style_font("Subtitle", 13, GRAY, False, "Calibri")
    style_font("Heading 1", 17, BLUE, True, "Calibri Light")
    style_font("Heading 2", 13.5, RED, True, "Calibri")
    style_font("Heading 3", 11.5, GRAY, True, "Calibri")
    style_font("Normal", 10.5, None, None, "Calibri")
    style_font("TOC Heading", 15, BLUE, True, "Calibri Light")

    # tiap bab (Heading 1) mulai di halaman baru
    try:
        d.styles["Heading 1"].paragraph_format.page_break_before = True
    except KeyError:
        pass
    for cs in ("Verbatim Char", "Source Code"):
        try:
            d.styles[cs].font.name = "Consolas"
            d.styles[cs].font.size = Pt(9)
        except KeyError:
            pass

    for para in d.paragraphs:
        if para.style.name == "Source Code":
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F3F5F7")
            pPr.append(shd)
            pbdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "16")
            left.set(qn("w:space"), "4"); left.set(qn("w:color"), "00519C")
            pbdr.append(left)
            pPr.append(pbdr)

    for tbl in d.tables:
        tblPr = tbl._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
            e.set(qn("w:color"), "B7C1CC")
            borders.append(e)
        tblPr.append(borders)
        if tbl.rows:
            for cell in tbl.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "00519C")
                tcPr.append(shd)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # gambar diagram: rata tengah + skala maksimal lebar area teks
    from docx.shared import Emu
    max_w = Emu(int((21.0 - 2 * 2.4) * 360000))  # A4 - margin, dalam EMU
    for para in d.paragraphs:
        drawings = para._p.findall(".//" + qn("w:drawing"))
        if drawings and not para.text.strip():
            para.alignment = 1  # center
            for shp in para._p.findall(".//" + qn("wp:extent")):
                cx, cy = int(shp.get("cx")), int(shp.get("cy"))
                if cx > max_w:
                    shp.set("cx", str(int(max_w)))
                    shp.set("cy", str(int(cy * (int(max_w) / cx))))

    # daftar isi = field TOC Word; minta Word meng-update semua field saat
    # dokumen dibuka (tanpa ini TOC tampil kosong)
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    d.settings.element.append(upd)

    for sec in d.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.2)
        sec.left_margin = sec.right_margin = Cm(2.4)

        # header: judul kecil abu-abu, rata kanan
        hp = sec.header.paragraphs[0]
        hp.text = TITLE
        hp.alignment = 2
        for r in hp.runs:
            r.font.size = Pt(8); r.font.color.rgb = GRAY

        # footer: "Halaman N" di tengah (field PAGE)
        fp = sec.footer.paragraphs[0]
        fp.text = ""
        fp.alignment = 1
        run = fp.add_run("Halaman ")
        run.font.size = Pt(8); run.font.color.rgb = GRAY
        fld_run = fp.add_run()
        fld_run.font.size = Pt(8); fld_run.font.color.rgb = GRAY
        for tag, attrs, txt in (
            ("w:fldChar", {"w:fldCharType": "begin"}, None),
            ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
            ("w:fldChar", {"w:fldCharType": "end"}, None),
        ):
            el = OxmlElement(tag)
            for k, v in attrs.items():
                el.set(qn(k), v)
            if txt:
                el.text = txt
            fld_run._r.append(el)

    d.save(str(OUT_DOCX))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", action="store_true", help="hasilkan juga .docx")
    args = ap.parse_args()

    mega = build_mega()
    OUT_MD.write_text(mega, encoding="utf-8", newline="\n")
    print(f"md  : {OUT_MD} ({OUT_MD.stat().st_size} bytes)")

    leftover = [m for m in re.finditer(r"\[([^\]]*\.md[^\]]*)\]", mega)
                if "bagian" not in m.group(1)]
    print(f"label masih berupa nama file: {len(leftover)}")

    if args.docx:
        md_to_docx(mega)
        print(f"docx: {OUT_DOCX} ({OUT_DOCX.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
